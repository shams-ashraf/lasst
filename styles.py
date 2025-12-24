import streamlit as st
import re
import fitz
import io
import docx
import uuid
import glob
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.utils import embedding_functions
import requests
import json
from datetime import datetime
from io import BytesIO
import os
import time
import pickle
import hashlib
from styles import load_custom_css


if 'chats' not in st.session_state:
    st.session_state.chats = {
        'default': {
            'messages': [],
            'current_context': [],
            'name': 'Chat 1'
        }
    }
    st.session_state.active_chat = 'default'

if 'processed' not in st.session_state:
    st.session_state.processed = False
    st.session_state.files_data = {}
    st.session_state.collection = None
    st.session_state.chat_counter = 1

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_MODEL = "llama-3.3-70b-versatile"
PDF_PASSWORD = "mbe2025"
DOCS_FOLDER = "/mount/src/lasst/documents"
CACHE_FOLDER = os.getenv("CACHE_FOLDER", "./cache")

os.makedirs(DOCS_FOLDER, exist_ok=True)
os.makedirs(CACHE_FOLDER, exist_ok=True)

if not GROQ_API_KEY:
    st.error("⚠️ GROQ_API_KEY not found in environment variables!")

def get_file_hash(filepath):
    hash_md5 = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def load_cache(cache_key):
    cache_file = os.path.join(CACHE_FOLDER, f"{cache_key}.pkl")
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'rb') as f:
                return pickle.load(f)
        except:
            return None
    return None

def save_cache(cache_key, data):
    cache_file = os.path.join(CACHE_FOLDER, f"{cache_key}.pkl")
    try:
        with open(cache_file, 'wb') as f:
            pickle.dump(data, f)
    except Exception as e:
        st.warning(f"⚠️ Could not save cache: {str(e)}")

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])
    return text.strip()

def structure_text_into_paragraphs(text):
    if not text or not text.strip():
        return ""
    
    text = clean_text(text)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    if not lines:
        return ""
    
    paragraphs = []
    current_paragraph = []
    
    for i, line in enumerate(lines):
        words_in_line = line.split()
        
        if len(words_in_line) < 3 and not (line[0].isupper() or re.match(r'^[\d]+[\.\):]', line)):
            continue
        
        is_heading = (
            (line.isupper() and len(words_in_line) <= 10) or
            (len(words_in_line) <= 6 and line[0].isupper() and line.endswith(':'))
        )
        
        if is_heading:
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                paragraph_text = re.sub(r'\s+', ' ', paragraph_text)
                paragraph_text = re.sub(r'\s+([.,!?;:])', r'\1', paragraph_text)
                paragraphs.append(paragraph_text.strip())
                current_paragraph = []
            paragraphs.append(f"\n🔹 {line}\n")
            continue
        
        is_list_item = re.match(r'^[\d]+[\.\)]\s', line) or re.match(r'^[•\-\*]\s', line)
        
        if is_list_item:
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                paragraph_text = re.sub(r'\s+', ' ', paragraph_text)
                paragraph_text = re.sub(r'\s+([.,!?;:])', r'\1', paragraph_text)
                paragraphs.append(paragraph_text.strip())
                current_paragraph = []
            paragraphs.append(f"  {line}")
            continue
        
        current_paragraph.append(line)
        
        ends_with_punctuation = line.endswith(('.', '!', '?', '؟', '!', '。'))
        next_is_new_section = False
        
        if i < len(lines) - 1:
            next_line = lines[i + 1]
            next_words = next_line.split()
            next_is_new_section = (
                re.match(r'^[\d]+[\.\)]\s', next_line) or
                re.match(r'^[•\-\*]\s', next_line) or
                (len(next_words) <= 6 and next_line[0].isupper()) or
                next_line.isupper()
            )
        
        is_last_line = (i == len(lines) - 1)
        
        if (ends_with_punctuation or next_is_new_section or is_last_line):
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                paragraph_text = re.sub(r'\s+', ' ', paragraph_text)
                paragraph_text = re.sub(r'\s+([.,!?;:])', r'\1', paragraph_text)
                paragraph_text = re.sub(r'([.,!?;:])\s*([.,!?;:])', r'\1', paragraph_text)
                paragraphs.append(paragraph_text.strip())
                current_paragraph = []
    
    if paragraphs:
        structured_text = ""
        for para in paragraphs:
            if para.startswith('\n🔹'):
                structured_text += para
            elif para.startswith('  '):
                structured_text += para + "\n"
            else:
                structured_text += para + "\n\n"
        return structured_text.strip()
    
    return text

def create_smart_chunks(text, chunk_size=1000, overlap=200, page_num=None, source_file=None, is_table=False, table_num=None):
    words = text.split()
    chunks = []
    
    metadata = {
        'page': str(page_num) if page_num is not None else "N/A",
        'source': source_file if source_file else "Unknown",
        'is_table': str(is_table),
        'table_number': str(table_num) if table_num is not None else "N/A"
    }
    
    if len(words) <= chunk_size:
        if text.strip():
            return [{
                'content': text,
                'metadata': metadata
            }]
        return []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk_words = words[i:i + chunk_size]
        chunk = " ".join(chunk_words)
        
        if len(chunk.split()) >= 30:
            chunks.append({
                'content': chunk,
                'metadata': metadata.copy()
            })
    
    return chunks

def format_table_as_structured_text(extracted_table, table_number=None):
    if not extracted_table or len(extracted_table) == 0:
        return ""
    
    headers = [str(cell).strip() if cell else f"Column_{i+1}" for i, cell in enumerate(extracted_table[0])]
    headers = [clean_text(h) for h in headers]
    
    text = f"\n📊 Table {table_number or ''}\n\n"
    
    if headers:
        text += "| " + " | ".join(headers) + " |\n"
        text += "| " + " --- |" * len(headers) + " |\n"
    
    row_count = 0
    for row in extracted_table[1:]:
        cells = [str(cell).strip() if cell else "" for cell in row]
        cells = [clean_text(c) for c in cells]
        
        if any(cells):
            text += "| " + " | ".join(cells) + " |\n"
            row_count += 1
    
    text += f"\n**Summary**: {row_count} data rows, {len(headers)} columns.\n"
    return text

def extract_pdf_detailed(filepath):
    try:
        doc = fitz.open(filepath)
        if doc.is_encrypted:
            if not doc.authenticate(PDF_PASSWORD):
                doc.close()
                return None, "❌ Invalid PDF password"
    except Exception as e:
        return None, f"❌ Error opening PDF: {str(e)}"
    
    filename = os.path.basename(filepath)
    file_info = {
        'chunks': [],
        'total_pages': len(doc),
        'total_tables': 0,
        'pages_with_tables': [],
    }
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        all_elements = []
        
        text_blocks = page.get_text("dict")["blocks"]
        for block in text_blocks:
            if block.get('type') == 0:
                y_pos = block.get('bbox', [0, 0, 0, 0])[1]
                text_content = ""
                for line in block.get('lines', []):
                    for span in line.get('spans', []):
                        text_content += span.get('text', '') + ' '
                
                if text_content.strip():
                    structured_content = structure_text_into_paragraphs(text_content)
                    all_elements.append({
                        'type': 'text',
                        'y_position': y_pos,
                        'content': structured_content,
                        'page': page_num + 1
                    })
        
        tables = page.find_tables()
        if tables and len(tables.tables) > 0:
            file_info['pages_with_tables'].append(page_num + 1)
            
            for table_num, table in enumerate(tables.tables, 1):
                file_info['total_tables'] += 1
                extracted_table = table.extract()
                
                if extracted_table:
                    table_text = format_table_as_structured_text(extracted_table, file_info['total_tables'])
                    all_elements.append({
                        'type': 'table',
                        'y_position': table.bbox[1] if table.bbox else 0,
                        'content': table_text,
                        'page': page_num + 1,
                        'table_num': file_info['total_tables']
                    })
        
        all_elements.sort(key=lambda x: x['y_position'])
        
        page_text = f"\n# Document: {filename} - Official MBE Regulations\n\n"
        page_text += f"\n{'═' * 60}\n📄 Page {page_num + 1}\n{'═' * 60}\n\n"
        
        for element in all_elements:
            page_text += element['content'] + "\n\n"
        
        page_chunks = create_smart_chunks(
            page_text,
            chunk_size=1500,
            overlap=250,
            page_num=page_num + 1,
            source_file=filename,
            is_table=False
        )
        
        for element in all_elements:
            if element['type'] == 'table':
                table_chunks = create_smart_chunks(
                    element['content'],
                    chunk_size=2000,
                    overlap=0,
                    page_num=element['page'],
                    source_file=filename,
                    is_table=True,
                    table_num=element.get('table_num')
                )
                file_info['chunks'].extend(table_chunks)
        
        file_info['chunks'].extend(page_chunks)
    
    doc.close()
    return file_info, None

def extract_docx_detailed(filepath):
    doc = docx.Document(filepath)
    filename = os.path.basename(filepath)
    
    file_info = {
        'chunks': [],
        'total_pages': 1,
        'total_tables': 0,
        'pages_with_tables': [],
    }
    
    all_text = []
    table_counter = 0
    
    for element in doc.element.body:
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element:
                    text = clean_text(para.text)
                    if text:
                        structured = structure_text_into_paragraphs(text)
                        if structured:
                            all_text.append(structured)
                    break
        
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    file_info['total_tables'] += 1
                    table_counter += 1
                    table_text = format_table_as_structured_text(
                        [[cell.text for cell in row.cells] for row in table.rows],
                        table_counter
                    )
                    
                    if table_text:
                        all_text.append(table_text)
                        
                        table_chunks = create_smart_chunks(
                            table_text,
                            chunk_size=2000,
                            overlap=0,
                            page_num=1,
                            source_file=filename,
                            is_table=True,
                            table_num=table_counter
                        )
                        file_info['chunks'].extend(table_chunks)
                    break
    
    complete_text = "\n\n".join(all_text)
    text_chunks = create_smart_chunks(
        complete_text,
        chunk_size=1500,
        overlap=250,
        page_num=1,
        source_file=filename
    )
    file_info['chunks'].extend(text_chunks)
    
    if file_info['total_tables'] > 0:
        file_info['pages_with_tables'] = [1]
    
    return file_info, None

def extract_txt_detailed(filepath):
    filename = os.path.basename(filepath)
    
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    
    structured_text = structure_text_into_paragraphs(text)
    chunks = create_smart_chunks(
        structured_text,
        chunk_size=1500,
        overlap=250,
        page_num=1,
        source_file=filename
    )
    
    file_info = {
        'chunks': chunks,
        'total_pages': 1,
        'total_tables': 0,
        'pages_with_tables': [],
    }
    
    return file_info, None

def get_files_from_folder():
    supported_extensions = ['*.pdf', '*.docx', '*.doc', '*.txt']
    files = []
    for ext in supported_extensions:
        files.extend(glob.glob(os.path.join(DOCS_FOLDER, ext)))
    return files

def get_embedding_function():
    return embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name="intfloat/multilingual-e5-large"
    )

def process_documents():
    available_files = get_files_from_folder()
    
    if not available_files:
        return False, "No documents found"
    
    files_data = {}
    all_chunks = []
    all_metadata = []
    
    client = chromadb.Client()
    collection_name = f"docs_{uuid.uuid4().hex[:8]}"
    collection = client.create_collection(
        name=collection_name,
        embedding_function=get_embedding_function()
    )
    
    for filepath in available_files:
        filename = os.path.basename(filepath)
        file_ext = filename.split('.')[-1].lower()
        
        file_hash = get_file_hash(filepath)
        cache_key = f"{file_hash}_{file_ext}"
        cached_data = load_cache(cache_key)
        
        if cached_data:
            file_info = cached_data
        else:
            if file_ext == 'pdf':
                file_info, error = extract_pdf_detailed(filepath)
            elif file_ext in ['docx', 'doc']:
                file_info, error = extract_docx_detailed(filepath)
            elif file_ext == 'txt':
                file_info, error = extract_txt_detailed(filepath)
            else:
                continue
            
            if error or not file_info:
                continue
            
            save_cache(cache_key, file_info)
        
        files_data[filename] = file_info
        
        for chunk_obj in file_info['chunks']:
            if isinstance(chunk_obj, dict):
                all_chunks.append(chunk_obj['content'])
                all_metadata.append(chunk_obj['metadata'])
            else:
                all_chunks.append(chunk_obj)
                all_metadata.append({
                    "source": filename,
                    "page": "N/A",
                    "is_table": "False",
                    "table_number": "N/A"
                })
    
    if all_chunks:
        batch_size = 500
        for i in range(0, len(all_chunks), batch_size):
            batch = all_chunks[i:i+batch_size]
            metadata_batch = all_metadata[i:i+batch_size]
            collection.add(
                documents=batch,
                ids=[f"chunk_{i+j}" for j in range(len(batch))],
                metadatas=metadata_batch
            )
        
        st.session_state.files_data = files_data
        st.session_state.collection = collection
        st.session_state.processed = True
        return True, len(available_files)
    
    return False, "No valid chunks extracted"

def answer_question_with_groq(query, relevant_chunks, chat_history=None):
    if not GROQ_API_KEY:
        return "❌ Please set GROQ_API_KEY in environment variables"
    
    context_parts = []
    for i, chunk_data in enumerate(relevant_chunks[:10], 1):
        content = chunk_data['content']
        meta = chunk_data['metadata']
        
        source = meta.get('source', 'Unknown')
        page = meta.get('page', 'N/A')
        is_table = meta.get('is_table', 'False')
        table_num = meta.get('table_number', 'N/A')
        
        citation = f"[Source {i}: {source}, Page {page}"
        if is_table == 'True' or is_table == True:
            citation += f", Table {table_num}"
        citation += "]"
        
        context_parts.append(f"{citation}\n{content}")
    
    context = "\n\n---\n\n".join(context_parts)
    
    conversation_summary = ""
    if chat_history and len(chat_history) > 0:
        recent = chat_history[-6:]
        conv_lines = []
        for msg in recent:
            role = "User" if msg['role'] == 'user' else "Assistant"
            content_preview = msg['content'][:300]
            conv_lines.append(f"{role}: {content_preview}")
        conversation_summary = "\n".join(conv_lines)
    
    data = {
        "model": GROQ_MODEL,
        "messages": [
            {
                "role": "system",
                "content": """You are a precise MBE Document Assistant at Hochschule Anhalt specializing in Biomedical Engineering regulations.

CRITICAL RULES:
1. Answer ONLY from provided sources OR previous conversation if it's a follow-up question
2. ALWAYS cite sources: [Source X, Page Y] or [Source X, Page Y, Table Z]
3. For follow-up questions like "summarize", "tell me more", "explain that", or "what about that":
   - Check the conversation history FIRST
   - Summarize or expand on your PREVIOUS answer
   - Don't search for new information if the question refers to what you just said
4. If user says "summarize that" or "summarize it": Condense your LAST answer (from conversation history)
5. If no relevant info in sources OR history: "No sufficient information in the available documents"
6. Use the SAME language as the question (English/German/Arabic)
7. Be CONCISE - short, direct answers unless asked to elaborate
8. For counting questions: Count precisely and list all items with citations

FOLLOW-UP DETECTION:
- "that", "it", "this", "summarize", "tell me more", "elaborate", "explain further" → Use conversation history
- New factual questions → Use sources

Remember: You're helping MBE students understand their program requirements clearly and accurately."""
            },
            {
                "role": "user",
                "content": f"""CONVERSATION HISTORY (use for follow-up questions):
{conversation_summary if conversation_summary else "No previous conversation"}

DOCUMENT SOURCES (use for new factual questions):
{context}

CURRENT QUESTION: {query}

Instructions:
- If this is a follow-up (summarize/elaborate/that/it), answer from conversation history
- If this is a new question, answer from sources with citations
- Always be precise and cite your sources

ANSWER:"""
            }
        ],
        "temperature": 0.1,
        "max_tokens": 2000,
    }
    
    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json=data,
            timeout=60
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"❌ Error connecting to Groq: {str(e)}"

if not st.session_state.processed:
    with st.spinner("🚀 Loading documents..."):
        success, result = process_documents()
        if success:
            st.success(f"✅ Loaded {result} document(s) successfully!")
        else:
            st.warning(f"⚠️ {result}")

st.markdown("""
<div style='text-align: center; padding: 1rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 10px; margin-bottom: 2rem;'>
    <h1 style='color: white; margin: 0;'>🎓 MBE Document Assistant</h1>
    <p style='color: white; margin: 0.5rem 0 0 0;'>RAG Chatbot for Biomedical Engineering at Hochschule Anhalt</p>
</div>
""", unsafe_allow_html=True)

with st.sidebar:
    st.markdown("### 💬 Chat Sessions")
    
    if st.button("➕ New Chat", use_container_width=True, type="primary"):
        st.session_state.chat_counter += 1
        new_chat_id = f"chat_{uuid.uuid4().hex[:8]}"
        st.session_state.chats[new_chat_id] = {
            'messages': [],
            'current_context': [],
            'name': f'Chat {st.session_state.chat_counter}'
        }
        st.session_state.active_chat = new_chat_id
        st.rerun()
    
    st.markdown("---")
    
    for chat_id, chat_data in list(st.session_state.chats.items()):
        col1, col2 = st.columns([4, 1])
        
        with col1:
            is_active = chat_id == st.session_state.active_chat
            button_type = "primary" if is_active else "secondary"
            if st.button(
                f"{'📌' if is_active else '💬'} {chat_data['name']}", 
                key=f"chat_{chat_id}",
                use_container_width=True,
                type=button_type
            ):
                st.session_state.active_chat = chat_id
                st.rerun()
        
        with col2:
            if st.button("🗑️", key=f"del_{chat_id}", help="Delete chat"):
                if len(st.session_state.chats) > 1:
                    del st.session_state.chats[chat_id]
                    if st.session_state.active_chat == chat_id:
                        st.session_state.active_chat = list(st.session_state.chats.keys())[0]
                    st.rerun()
                else:
                    st.warning("⚠️ Cannot delete the last chat!")
    
    st.markdown("---")
    
    if st.button("🗑️ Clear Cache & Reload", use_container_width=True):
        import shutil
        if os.path.exists(CACHE_FOLDER):
            shutil.rmtree(CACHE_FOLDER)
            os.makedirs(CACHE_FOLDER)
        st.session_state.processed = False
        st.session_state.files_data = {}
        st.session_state.collection = None
        st.success("✅ Cache cleared! Reloading...")
        st.rerun()

if st.session_state.processed and st.session_state.active_chat in st.session_state.chats:
    current_chat = st.session_state.chats[st.session_state.active_chat]
    
    st.markdown(f"### 💬 {current_chat['name']}")
    
    for message in current_chat['messages']:
        role = message["role"]
        content = message["content"]
        
        if role == "user":
            st.markdown(f'<div class="chat-message user-message">👤 <strong>You:</strong> {content}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="chat-message assistant-message">🤖 <strong>Assistant:</strong> {content}</div>', unsafe_allow_html=True)
    
    query = st.chat_input("Ask anything about your documents...")
    
    if query:
        current_chat['messages'].append({"role": "user", "content": query})
        
        with st.spinner("Thinking..."):
            results = st.session_state.collection.query(
                query_texts=[query],
                n_results=10
            )
            
            relevant_chunks = []
            for content, metadata in zip(results["documents"][0], results["metadatas"][0]):
                relevant_chunks.append({
                    'content': content,
                    'metadata': metadata
                })
            
            answer = answer_question_with_groq(query, relevant_chunks, current_chat['messages'])
            
            current_chat['messages'].append({"role": "assistant", "content": answer})
            
            current_chat['current_context'] = relevant_chunks
        
        st.rerun()
    
    if current_chat['current_context']:
        with st.expander("📄 View Sources", expanded=False):
            for idx, chunk_data in enumerate(current_chat['current_context'][:5], 1):
                meta = chunk_data['metadata']
                source = meta.get('source', 'Unknown')
                page = meta.get('page', 'N/A')
                is_table = meta.get('is_table', 'False')
                table_num = meta.get('table_number', 'N/A')
                
                citation_info = f"📄 **Source {idx}**: {source} | Page {page}"
                if is_table == 'True' or is_table == True:
                    citation_info += f" | Table {table_num}"
                
                st.markdown(citation_info)
                st.markdown(f'<div class="source-box">{chunk_data["content"][:500]}...</div>', unsafe_allow_html=True)
                st.markdown("---")

else:
    st.info("📁 Documents are loading or not found. Please check the documents folder.")

st.markdown("""
---
### 🎯 Features:
- **Multi-Chat Support**: Create multiple chat sessions for different topics
- **Precise Citations**: Every answer includes file + page + table references
- **Conversational**: Ask follow-up questions naturally
- **Auto-Load**: Documents process automatically on startup
- **Fast**: Cached processing for instant responses
    ### 📋 Supported Documents:
    - 📄 Study & Examination Regulations (SPO)
    - 📚 Module Handbook
    - 📝 Guide for Writing Scientific Papers
    - 📃 Notes on Bachelor/Master Theses
    - ✍️ Scientific Writing Guidelines
    
    ### 💡 Example Questions:
    - "How many modules in semester 1?"
    - "What are the thesis requirements?"
    - "Tell me about the internship" → then "summarize that"
    - "Compare exam types in SPO"
    """)
