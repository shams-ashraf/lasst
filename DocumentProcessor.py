import streamlit as st
import re
import fitz  # PyMuPDF
import docx
import glob
import os
import pickle
import hashlib
from datetime import datetime

PDF_PASSWORD = os.getenv("PDF_PASSWORD", "")
DOCS_FOLDER = "/mount/src/lasst/documents"
CACHE_FOLDER = os.getenv("CACHE_FOLDER", "./cache")

os.makedirs(DOCS_FOLDER, exist_ok=True)
os.makedirs(CACHE_FOLDER, exist_ok=True)

# ========== تحسين 1: Cache مع timestamp ==========
def get_file_hash(filepath):
    """Hash للملف + تاريخ التعديل"""
    hash_md5 = hashlib.md5()
    
    # Hash محتوى الملف
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    
    # إضافة timestamp
    mtime = os.path.getmtime(filepath)
    hash_md5.update(str(mtime).encode())
    
    return hash_md5.hexdigest()

def load_cache(cache_key):
    cache_file = os.path.join(CACHE_FOLDER, f"{cache_key}.pkl")
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'rb') as f:
                cached_data = pickle.load(f)
                
                # التحقق من صلاحية الكاش (30 يوم)
                if 'timestamp' in cached_data:
                    cache_age = datetime.now() - cached_data['timestamp']
                    if cache_age.days > 30:
                        return None  # كاش قديم
                
                return cached_data.get('data')
        except:
            return None
    return None

def save_cache(cache_key, data):
    cache_file = os.path.join(CACHE_FOLDER, f"{cache_key}.pkl")
    try:
        cache_data = {
            'data': data,
            'timestamp': datetime.now()
        }
        with open(cache_file, 'wb') as f:
            pickle.dump(cache_data, f)
    except Exception as e:
        st.warning(f"⚠️ Cache save error: {str(e)}")

# ========== تحسين 2: تنظيف النصوص أفضل ==========
def clean_text(text):
    """تنظيف متقدم للنصوص"""
    # إزالة أحرف غريبة
    text = re.sub(r'[^\w\s\.\,\:\;\!\?\-\(\)\[\]\{\}\"\'\/\n\r\u0600-\u06FF]', '', text)
    
    # توحيد المسافات
    text = re.sub(r'\s+', ' ', text).strip()
    
    # إزالة أرقام الصفحات المعزولة
    text = re.sub(r'\b\d{1,3}\b(?=\s|$)', '', text)
    
    return text

def structure_text_into_paragraphs(text):
    """تنظيم النص في فقرات منطقية"""
    if not text.strip():
        return ""
    
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    paragraphs = []
    current = []
    
    for line in lines:
        # كشف العناوين والقوائم
        is_heading = re.match(r'^[\d]+[\.\)]\s|^[•\-\*]\s|^#{1,3}\s|^[A-Z\u0600-\u06FF][A-Z\u0600-\u06FF\s]{3,}$', line)
        
        if is_heading:
            if current:
                paragraphs.append(' '.join(current))
                current = []
            paragraphs.append(line)
        else:
            # دمج السطور القصيرة
            if len(line) < 50 and current:
                current[-1] += ' ' + line
            else:
                current.append(line)
    
    if current:
        paragraphs.append(' '.join(current))
    
    return '\n\n'.join(paragraphs)

# ========== تحسين 3: Chunking أذكى ==========
def create_smart_chunks(text, chunk_size=800, overlap=100, page_num=None, source_file=None, is_table=False, table_num=None):
    """تقسيم ذكي مع الحفاظ على السياق"""
    
    # الجداول: لا نقسمها إلا إذا كانت ضخمة
    if is_table and len(text.split()) < 500:
        metadata = {
            'page': str(page_num) if page_num is not None else "N/A",
            'source': source_file or "Unknown",
            'is_table': 'True',
            'table_number': str(table_num) if table_num else "N/A",
            'chunk_type': 'complete_table'
        }
        return [{'content': text.strip(), 'metadata': metadata}]
    
    # تقسيم حسب الفقرات أولاً
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = []
    current_size = 0
    
    for para in paragraphs:
        para_words = para.split()
        para_size = len(para_words)
        
        # إذا الفقرة وحدها أكبر من chunk_size
        if para_size > chunk_size:
            # حفظ current chunk
            if current_chunk:
                chunk_text = ' '.join(current_chunk)
                metadata = {
                    'page': str(page_num) if page_num is not None else "N/A",
                    'source': source_file or "Unknown",
                    'is_table': str(is_table),
                    'table_number': str(table_num) if table_num else "N/A",
                    'chunk_type': 'paragraph_group'
                }
                chunks.append({'content': chunk_text, 'metadata': metadata})
                current_chunk = []
                current_size = 0
            
            # تقسيم الفقرة الكبيرة
            for i in range(0, para_size, chunk_size - overlap):
                sub_chunk = para_words[i:i + chunk_size]
                if len(sub_chunk) >= 50:  # تجاهل القطع الصغيرة جداً
                    metadata = {
                        'page': str(page_num) if page_num is not None else "N/A",
                        'source': source_file or "Unknown",
                        'is_table': str(is_table),
                        'chunk_type': 'large_paragraph_split'
                    }
                    chunks.append({'content': ' '.join(sub_chunk), 'metadata': metadata})
        
        # إضافة فقرة عادية
        elif current_size + para_size <= chunk_size:
            current_chunk.append(para)
            current_size += para_size
        
        # الفقرة لا تتسع في current chunk
        else:
            if current_chunk:
                chunk_text = '\n\n'.join(current_chunk)
                metadata = {
                    'page': str(page_num) if page_num is not None else "N/A",
                    'source': source_file or "Unknown",
                    'is_table': str(is_table),
                    'chunk_type': 'paragraph_group'
                }
                chunks.append({'content': chunk_text, 'metadata': metadata})
            
            current_chunk = [para]
            current_size = para_size
    
    # آخر chunk
    if current_chunk:
        chunk_text = '\n\n'.join(current_chunk)
        metadata = {
            'page': str(page_num) if page_num is not None else "N/A",
            'source': source_file or "Unknown",
            'is_table': str(is_table),
            'chunk_type': 'paragraph_group'
        }
        chunks.append({'content': chunk_text, 'metadata': metadata})
    
    return chunks

# ========== تحسين 4: معالجة الجداول أفضل ==========
def format_table_as_structured_text(table, table_number=None):
    """تحويل الجدول لنص منظم مع معالجة الخلايا المدمجة"""
    if not table or len(table) == 0:
        return ""
    
    # تنظيف البيانات
    clean_table = []
    for row in table:
        clean_row = [str(cell).strip() if cell else "" for cell in row]
        clean_table.append(clean_row)
    
    # العناوين
    headers = clean_table[0] if clean_table else []
    headers = [h or f"Col_{i+1}" for i, h in enumerate(headers)]
    
    # بناء النص
    text = f"\n📊 **Table {table_number or ''}**\n\n"
    text += "| " + " | ".join(headers) + " |\n"
    text += "|" + "---|" * len(headers) + "\n"
    
    # البيانات
    for row in clean_table[1:]:
        # تجاهل الصفوف الفارغة
        if not any(cell.strip() for cell in row):
            continue
        
        # تعبئة الخلايا الناقصة
        while len(row) < len(headers):
            row.append("")
        
        text += "| " + " | ".join(row[:len(headers)]) + " |\n"
    
    return text

# ========== باقي الدوال (بدون تغيير كبير) ==========
def extract_pdf_detailed(filepath):
    try:
        doc = fitz.open(filepath)
        if doc.is_encrypted and not doc.authenticate(PDF_PASSWORD):
            return None, "❌ Wrong PDF password"
    except Exception as e:
        return None, f"❌ PDF open error: {str(e)}"

    filename = os.path.basename(filepath)
    file_info = {'chunks': [], 'total_pages': len(doc), 'total_tables': 0}

    for page_num in range(len(doc)):
        page = doc[page_num]

        # OCR إذا لزم الأمر
        text = page.get_text("text")
        if len(text.strip()) < 100:
            textpage = page.get_textpage_ocr(
                flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE, 
                full=True
            )
            text = page.get_text("text", textpage=textpage)

        # استخراج النص المنظم
        blocks = page.get_text("dict")["blocks"]
        page_text = f"# {filename} - Page {page_num + 1}\n\n"
        
        for block in blocks:
            if block.get("type") == 0:  # نص
                block_text = ""
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        block_text += span.get("text", "")
                
                if block_text.strip():
                    cleaned = clean_text(block_text)
                    structured = structure_text_into_paragraphs(cleaned)
                    page_text += structured + "\n\n"

        # استخراج الجداول
        tables = page.find_tables()
        if tables:
            for t_num, table in enumerate(tables.tables, 1):
                file_info['total_tables'] += 1
                extracted = table.extract()
                
                if extracted:
                    table_text = format_table_as_structured_text(
                        extracted, 
                        file_info['total_tables']
                    )
                    page_text += table_text + "\n\n"
                    
                    # Chunk خاص بالجدول
                    table_chunks = create_smart_chunks(
                        table_text, 
                        chunk_size=2000,  # جداول أكبر
                        overlap=0, 
                        page_num=page_num+1,
                        source_file=filename, 
                        is_table=True, 
                        table_num=file_info['total_tables']
                    )
                    file_info['chunks'].extend(table_chunks)

        # Chunks للنص العادي
        page_chunks = create_smart_chunks(
            page_text, 
            chunk_size=1000,  # زيادة قليلة
            overlap=150, 
            page_num=page_num+1, 
            source_file=filename
        )
        file_info['chunks'].extend(page_chunks)

    doc.close()
    return file_info, None

def extract_docx_detailed(filepath):
    """معالجة DOCX محسّنة"""
    doc = docx.Document(filepath)
    filename = os.path.basename(filepath)
    file_info = {
        'chunks': [],
        'total_pages': 1,
        'total_tables': 0,
    }
   
    all_text = []
    table_counter = 0
   
    for element in doc.element.body:
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element:
                    text = clean_text(para.text)
                    if text:
                        structured = structure_text_into_paragraphs(text)
                        if structured:
                            all_text.append(structured)
                    break
       
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    file_info['total_tables'] += 1
                    table_counter += 1
                    
                    table_data = [[cell.text for cell in row.cells] for row in table.rows]
                    table_text = format_table_as_structured_text(
                        table_data,
                        table_counter
                    )
                    
                    if table_text:
                        all_text.append(table_text)
                        
                        table_chunks = create_smart_chunks(
                            table_text,
                            chunk_size=2000,
                            overlap=0,
                            page_num=1,
                            source_file=filename,
                            is_table=True,
                            table_num=table_counter
                        )
                        file_info['chunks'].extend(table_chunks)
                    break
   
    complete_text = "\n\n".join(all_text)
    text_chunks = create_smart_chunks(
        complete_text, 
        chunk_size=1200, 
        overlap=200,
        page_num=1,
        source_file=filename
    )
    file_info['chunks'].extend(text_chunks)
   
    return file_info, None

def extract_txt_detailed(filepath):
    filename = os.path.basename(filepath)
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    
    cleaned_text = clean_text(text)
    structured_text = structure_text_into_paragraphs(cleaned_text)
    
    chunks = create_smart_chunks(
        structured_text, 
        chunk_size=1200, 
        overlap=200,
        page_num=1,
        source_file=filename
    )
    
    file_info = {
        'chunks': chunks,
        'total_pages': 1,
        'total_tables': 0,
    }
    return file_info, None
    
def get_files_from_folder():
    """جلب جميع الملفات المدعومة"""
    patterns = [
        "*.[pP][dD][fF]",
        "*.[dD][oO][cC][xX]",
        "*.txt"
    ]
    
    files = []
    for pattern in patterns:
        files.extend(glob.glob(os.path.join(DOCS_FOLDER, pattern)))
    
    return sorted(files)  # ترتيب أبجدي
