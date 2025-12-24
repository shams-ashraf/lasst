import streamlit as st
import re
import fitz
import docx
import uuid
import glob
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.utils import embedding_functions
import requests
import os
import pickle
import hashlib

# Page Configuration
st.set_page_config(
    page_title="MBE Document Assistant",
    page_icon="🎓",
    layout="wide"
)

# Modern Dark UI Styling
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');

* {font-family: 'Inter', sans-serif;}

.main {
    background: linear-gradient(135deg, #0a0a0a 0%, #1a1a2e 50%, #0f0f0f 100%);
    color: #e8e8e8;
}

[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #16213e 0%, #0f3460 100%);
    padding: 1.5rem 1rem;
    box-shadow: 4px 0 20px rgba(0, 0, 0, 0.5);
}

[data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
    color: #00d9ff !important;
    text-shadow: 0 0 15px rgba(0, 217, 255, 0.4);
    font-weight: 700;
    margin-bottom: 0.5rem !important;
}

[data-testid="stSidebar"] p, [data-testid="stSidebar"] span, [data-testid="stSidebar"] label {
    color: #ffffff !important;
}

.stButton button {
    background: linear-gradient(90deg, #00d9ff 0%, #0099cc 100%);
    color: white;
    border: none;
    border-radius: 12px;
    padding: 0.7rem 1.5rem;
    font-weight: 600;
    transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
    box-shadow: 0 4px 20px rgba(0, 217, 255, 0.3);
    width: 100%;
    margin: 8px 0;
    font-size: 0.95rem;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}

.stButton button:hover {
    transform: translateY(-3px) scale(1.02);
    box-shadow: 0 8px 30px rgba(0, 217, 255, 0.6);
    background: linear-gradient(90deg, #00e5ff 0%, #00b8e6 100%);
}

.chat-message {
    padding: 1.5rem;
    border-radius: 18px;
    margin-bottom: 1.5rem;
    animation: slideIn 0.6s cubic-bezier(0.4, 0, 0.2, 1);
    box-shadow: 0 4px 20px rgba(0, 0, 0, 0.08);
}

@keyframes slideIn {
    from {
        opacity: 0;
        transform: translateY(20px);
    }
    to {
        opacity: 1;
        transform: translateY(0);
    }
}

.user-message {
    background: linear-gradient(135deg, #378ecf 0%, #07547e 100%);
    border-left: 5px solid #0c54b0;
}

.user-message div:last-child {
    color: #ffffff !important;
    font-weight: 500;
}

.assistant-message {
    background: linear-gradient(135deg, #2d3748 0%, #1a202c 100%);
    border-left: 5px solid #48bb78;
    color: #e2e8f0;
}

.message-header {
    font-weight: 700;
    margin-bottom: 10px;
    font-size: 0.85rem;
    text-transform: uppercase;
    letter-spacing: 1.5px;
}

.user-message .message-header {
    color: #00d9ff;
}

.assistant-message .message-header {
    color: #48bb78;
}

.stTextInput input {
    background: rgb(10 144 168 / 50%);
    border: 2px solid #0ba4be;
    border-radius: 15px;
    color: #ffffff;
    padding: 1.5rem 1.2rem;
    font-size: 1.15rem;
    transition: all 0.3s;
}

.stTextInput input::placeholder {
    color: #ffffff;
    opacity: 0.8;
}

.stTextInput input:focus {
    border-color: #00e5ff;
    box-shadow: 0 0 25px rgba(0, 217, 255, 0.7);
    background: rgba(0, 217, 255, 0.6);
}

.file-badge {
    display: inline-block;
    background: linear-gradient(90deg, #10b981 0%, #059669 100%);
    color: white;
    padding: 8px 18px;
    border-radius: 25px;
    margin: 6px 4px;
    font-size: 0.88rem;
    font-weight: 600;
    box-shadow: 0 4px 15px rgba(16, 185, 129, 0.3);
    transition: all 0.3s;
}

.file-badge:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 20px rgba(16, 185, 129, 0.5);
}

h1 {
    color: #0f63bc !important;
    text-align: center;
    font-size: 3rem !important;
    text-shadow: 0 0 30px rgba(0, 217, 255, 0.6);
    margin-bottom: 2.5rem !important;
    font-weight: 800 !important;
    letter-spacing: -1px;
}

.stInfo {
    background: linear-gradient(135deg, rgba(0, 217, 255, 0.15) 0%, rgba(0, 153, 204, 0.15) 100%);
    border-left: 5px solid #00d9ff;
    border-radius: 12px;
    padding: 1rem;
    color: #e8e8e8;
}

.stSuccess {
    background: linear-gradient(135deg, rgba(72, 187, 120, 0.15) 0%, rgba(56, 161, 105, 0.15) 100%);
    border-left: 5px solid #48bb78;
    border-radius: 12px;
    color: #e8e8e8;
}

.stWarning {
    background: linear-gradient(135deg, rgba(237, 137, 54, 0.15) 0%, rgba(221, 107, 32, 0.15) 100%);
    border-left: 5px solid #ed8936;
    border-radius: 12px;
    color: #e8e8e8;
}

.stError {
    background: linear-gradient(135deg, rgba(239, 68, 68, 0.15) 0%, rgba(220, 38, 38, 0.15) 100%);
    border-left: 5px solid #ef4444;
    border-radius: 12px;
    color: #e8e8e8;
}

::-webkit-scrollbar {width: 12px;}
::-webkit-scrollbar-track {background: #1a1a2e;}
::-webkit-scrollbar-thumb {
    background: linear-gradient(180deg, #ffffff 0%, #e2e8f0 100%);
    border-radius: 10px;
    border: 2px solid #00d9ff;
}
::-webkit-scrollbar-thumb:hover {
    background: linear-gradient(180deg, #e2e8f0 0%, #cbd5e1 100%);
}

.stat-card {
    background: linear-gradient(135deg, rgba(102, 126, 234, 0.2) 0%, rgba(118, 75, 162, 0.2) 100%);
    border-radius: 15px;
    padding: 1.5rem;
    margin: 1rem 0;
    border-left: 4px solid #667eea;
    box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3);
}

.stat-card h3 {
    color: #00d9ff;
    margin-bottom: 0.5rem;
}

.stat-card p {
    color: #e8e8e8;
    font-size: 0.95rem;
}
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'processed' not in st.session_state:
    st.session_state.processed = False
    st.session_state.files_data = {}
    st.session_state.collection = None
    st.session_state.messages = []
    st.session_state.current_context = []
    st.session_state.processing_started = False

# Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_MODEL = "llama-3.3-70b-versatile"
PDF_PASSWORD = "mbe2025"
DOCS_FOLDER = "/mount/src/test/documents"
CACHE_FOLDER = os.getenv("CACHE_FOLDER", "./cache")

os.makedirs(DOCS_FOLDER, exist_ok=True)
os.makedirs(CACHE_FOLDER, exist_ok=True)

# Helper Functions
def get_file_hash(filepath):
    """Calculate MD5 hash of file for caching"""
    hash_md5 = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def load_cache(cache_key):
    """Load processed data from cache"""
    cache_file = os.path.join(CACHE_FOLDER, f"{cache_key}.pkl")
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'rb') as f:
                return pickle.load(f)
        except:
            return None
    return None

def save_cache(cache_key, data):
    """Save processed data to cache"""
    cache_file = os.path.join(CACHE_FOLDER, f"{cache_key}.pkl")
    try:
        with open(cache_file, 'wb') as f:
            pickle.dump(data, f)
    except Exception as e:
        st.warning(f"⚠️ Could not save cache: {str(e)}")

def clean_text(text):
    """Clean and normalize text"""
    text = re.sub(r'\s+', ' ', text)
    text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])
    return text.strip()

def structure_text_into_paragraphs(text):
    """Structure raw text into readable paragraphs"""
    if not text or not text.strip():
        return ""
   
    text = clean_text(text)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
   
    if not lines:
        return ""
   
    paragraphs = []
    current_paragraph = []
   
    for i, line in enumerate(lines):
        words_in_line = line.split()
       
        if len(words_in_line) < 3 and not (line[0].isupper() or re.match(r'^[\d]+[\.\):]', line)):
            continue
       
        is_heading = (
            (line.isupper() and len(words_in_line) <= 10) or
            (len(words_in_line) <= 6 and line[0].isupper() and line.endswith(':'))
        )
       
        if is_heading:
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                paragraph_text = re.sub(r'\s+', ' ', paragraph_text)
                paragraph_text = re.sub(r'\s+([.,!?;:])', r'\1', paragraph_text)
                paragraphs.append(paragraph_text.strip())
                current_paragraph = []
            paragraphs.append(f"\n🔹 {line}\n")
            continue
       
        is_list_item = re.match(r'^[\d]+[\.\)]\s', line) or re.match(r'^[•\-\*]\s', line)
       
        if is_list_item:
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                paragraph_text = re.sub(r'\s+', ' ', paragraph_text)
                paragraph_text = re.sub(r'\s+([.,!?;:])', r'\1', paragraph_text)
                paragraphs.append(paragraph_text.strip())
                current_paragraph = []
            paragraphs.append(f" {line}")
            continue
       
        current_paragraph.append(line)
       
        ends_with_punctuation = line.endswith(('.', '!', '?', '؟', '!', '。'))
        next_is_new_section = False
       
        if i < len(lines) - 1:
            next_line = lines[i + 1]
            next_words = next_line.split()
            next_is_new_section = (
                re.match(r'^[\d]+[\.\)]\s', next_line) or
                re.match(r'^[•\-\*]\s', next_line) or
                (len(next_words) <= 6 and next_line[0].isupper()) or
                next_line.isupper()
            )
       
        is_last_line = (i == len(lines) - 1)
       
        if (ends_with_punctuation or next_is_new_section or is_last_line):
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                paragraph_text = re.sub(r'\s+', ' ', paragraph_text)
                paragraph_text = re.sub(r'\s+([.,!?;:])', r'\1', paragraph_text)
                paragraph_text = re.sub(r'([.,!?;:])\s*([.,!?;:])', r'\1', paragraph_text)
                paragraphs.append(paragraph_text.strip())
                current_paragraph = []
   
    if paragraphs:
        structured_text = ""
        for para in paragraphs:
            if para.startswith('\n🔹'):
                structured_text += para
            elif para.startswith(' '):
                structured_text += para + "\n"
            else:
                structured_text += para + "\n\n"
        return structured_text.strip()
   
    return text

def create_smart_chunks(text, chunk_size=1000, overlap=200, page_num=None, source_file=None, is_table=False, table_num=None):
    """Create smart chunks with metadata for vector storage"""
    words = text.split()
    chunks = []
    
    metadata = {
        'page': str(page_num) if page_num is not None else "N/A",
        'source': source_file if source_file else "Unknown",
        'is_table': str(is_table),
        'table_number': str(table_num) if table_num is not None else "N/A"
    }
   
    if len(words) <= chunk_size:
        if text.strip():
            return [{
                'content': text,
                'metadata': metadata
            }]
        return []
   
    for i in range(0, len(words), chunk_size - overlap):
        chunk_words = words[i:i + chunk_size]
        chunk = " ".join(chunk_words)
        if len(chunk.split()) >= 30:
            chunks.append({
                'content': chunk,
                'metadata': metadata.copy()
            })
   
    return chunks

def format_table_as_structured_text(extracted_table, table_number=None):
    """Format table data as structured markdown text"""
    if not extracted_table or len(extracted_table) == 0:
        return ""
   
    headers = [str(cell).strip() if cell else f"Column_{i+1}" for i, cell in enumerate(extracted_table[0])]
    headers = [clean_text(h) for h in headers]
   
    text = f"\n📊 Table {table_number or ''}\n\n"
    if headers:
        text += "| " + " | ".join(headers) + " |\n"
        text += "| " + " --- |" * len(headers) + " |\n"
    
    row_count = 0
    for row in extracted_table[1:]:
        cells = [str(cell).strip() if cell else "" for cell in row]
        cells = [clean_text(c) for c in cells]
        if any(cells):
            text += "| " + " | ".join(cells) + " |\n"
            row_count += 1
    
    text += f"\n**Summary**: {row_count} data rows, {len(headers)} columns.\n"
    return text

def extract_pdf_detailed(filepath):
    """Extract text and tables from PDF with structure preservation"""
    try:
        doc = fitz.open(filepath)
        if doc.is_encrypted:
            if not doc.authenticate(PDF_PASSWORD):
                doc.close()
                return None, "❌ Invalid PDF password"
    except Exception as e:
        return None, f"❌ Error opening PDF: {str(e)}"
   
    filename = os.path.basename(filepath)
    file_info = {
        'chunks': [],
        'total_pages': len(doc),
        'total_tables': 0,
        'pages_with_tables': [],
    }
   
    for page_num in range(len(doc)):
        page = doc[page_num]
        all_elements = []
       
        text_blocks = page.get_text("dict")["blocks"]
        for block in text_blocks:
            if block.get('type') == 0:
                y_pos = block.get('bbox', [0, 0, 0, 0])[1]
                text_content = ""
                for line in block.get('lines', []):
                    for span in line.get('spans', []):
                        text_content += span.get('text', '') + ' '
                if text_content.strip():
                    structured_content = structure_text_into_paragraphs(text_content)
                    all_elements.append({
                        'type': 'text',
                        'y_position': y_pos,
                        'content': structured_content,
                        'page': page_num + 1
                    })
       
        tables = page.find_tables()
        if tables and len(tables.tables) > 0:
            file_info['pages_with_tables'].append(page_num + 1)
           
            for table_num, table in enumerate(tables.tables, 1):
                file_info['total_tables'] += 1
                extracted_table = table.extract()
                if extracted_table:
                    table_text = format_table_as_structured_text(extracted_table, file_info['total_tables'])
                    all_elements.append({
                        'type': 'table',
                        'y_position': table.bbox[1] if table.bbox else 0,
                        'content': table_text,
                        'page': page_num + 1,
                        'table_num': file_info['total_tables']
                    })
       
        all_elements.sort(key=lambda x: x['y_position'])
       
        page_text = f"\n# Document: {filename} - Official MBE Regulations\n\n"
        page_text += f"\n{'═' * 60}\n📄 Page {page_num + 1}\n{'═' * 60}\n\n"
        for element in all_elements:
            page_text += element['content'] + "\n\n"
       
        page_chunks = create_smart_chunks(
            page_text, 
            chunk_size=1500, 
            overlap=250,
            page_num=page_num + 1,
            source_file=filename,
            is_table=False
        )
        
        for element in all_elements:
            if element['type'] == 'table':
                table_chunks = create_smart_chunks(
                    element['content'],
                    chunk_size=2000,
                    overlap=0,
                    page_num=element['page'],
                    source_file=filename,
                    is_table=True,
                    table_num=element.get('table_num')
                )
                file_info['chunks'].extend(table_chunks)
        
        file_info['chunks'].extend(page_chunks)
   
    doc.close()
    return file_info, None

def extract_docx_detailed(filepath):
    """Extract text and tables from DOCX files"""
    doc = docx.Document(filepath)
    filename = os.path.basename(filepath)
    file_info = {
        'chunks': [],
        'total_pages': 1,
        'total_tables': 0,
        'pages_with_tables': [],
    }
   
    all_text = []
    table_counter = 0
   
    for element in doc.element.body:
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element:
                    text = clean_text(para.text)
                    if text:
                        structured = structure_text_into_paragraphs(text)
                        if structured:
                            all_text.append(structured)
                    break
       
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    file_info['total_tables'] += 1
                    table_counter += 1
                    table_text = format_table_as_structured_text(
                        [[cell.text for cell in row.cells] for row in table.rows],
                        table_counter
                    )
                    if table_text:
                        all_text.append(table_text)
                        table_chunks = create_smart_chunks(
                            table_text,
                            chunk_size=2000,
                            overlap=0,
                            page_num=1,
                            source_file=filename,
                            is_table=True,
                            table_num=table_counter
                        )
                        file_info['chunks'].extend(table_chunks)
                    break
   
    complete_text = "\n\n".join(all_text)
    text_chunks = create_smart_chunks(
        complete_text, 
        chunk_size=1500, 
        overlap=250,
        page_num=1,
        source_file=filename
    )
    file_info['chunks'].extend(text_chunks)
   
    if file_info['total_tables'] > 0:
        file_info['pages_with_tables'] = [1]
   
    return file_info, None

def extract_txt_detailed(filepath):
    """Extract text from TXT files"""
    filename = os.path.basename(filepath)
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    structured_text = structure_text_into_paragraphs(text)
    chunks = create_smart_chunks(
        structured_text, 
        chunk_size=1500, 
        overlap=250,
        page_num=1,
        source_file=filename
    )
    file_info = {
        'chunks': chunks,
        'total_pages': 1,
        'total_tables': 0,
        'pages_with_tables': [],
    }
    return file_info, None

def get_files_from_folder():
    """Get all supported document files from folder"""
    supported_extensions = ['*.pdf', '*.docx', '*.doc', '*.txt']
    files = []
    for ext in supported_extensions:
        files.extend(glob.glob(os.path.join(DOCS_FOLDER, ext)))
    return files

def get_embedding_function():
    """Get multilingual embedding function for vector storage"""
    return embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name="intfloat/multilingual-e5-large"
    )

def answer_question_with_groq(query, relevant_chunks, chat_history=None):
    """Generate answer using Groq LLM with context and chat history"""
    if not GROQ_API_KEY:
        return "❌ Please set GROQ_API_KEY in environment variables"
   
    context_parts = []
    for i, chunk_data in enumerate(relevant_chunks[:10], 1):
        content = chunk_data['content']
        meta = chunk_data['metadata']
        
        source = meta.get('source', 'Unknown')
        page = meta.get('page', 'N/A')
        is_table = meta.get('is_table', 'False')
        table_num = meta.get('table_number', 'N/A')
        
        citation = f"[Source {i}: {source}, Page {page}"
        if is_table == 'True' or is_table == True:
            citation += f", Table {table_num}"
        citation += "]"
        
        context_parts.append(f"{citation}\n{content}")
    
    context = "\n\n---\n\n".join(context_parts)
    
    conversation_summary = ""
    if chat_history and len(chat_history) > 0:
        recent = chat_history[-6:]
        conv_lines = []
        for msg in recent:
            role = "User" if msg['role'] == 'user' else "Assistant"
            content_preview = msg['content'][:300]
            conv_lines.append(f"{role}: {content_preview}")
        conversation_summary = "\n".join(conv_lines)
   
    data = {
        "model": GROQ_MODEL,
        "messages": [
            {
                "role": "system",
                "content": """You are a precise MBE Document Assistant at Hochschule Anhalt specializing in Biomedical Engineering regulations.

CRITICAL RULES:
1. Answer ONLY from provided sources OR previous conversation if it's a follow-up question
2. ALWAYS cite sources: [Source X, Page Y] or [Source X, Page Y, Table Z]
3. For follow-up questions like "summarize", "tell me more", "explain that", or "what about that":
   - Check the conversation history FIRST
   - Summarize or expand on your PREVIOUS answer
   - Don't search for new information if the question refers to what you just said
4. If user says "summarize that" or "summarize it": Condense your LAST answer (from conversation history)
5. If no relevant info in sources OR history: "No sufficient information in the available documents"
6. Use the SAME language as the question:
   - English question → English answer
   - German question → German answer (Deutsch)
7. Be CONCISE - short, direct answers unless asked to elaborate
8. For counting questions: Count precisely and list all items with citations

LANGUAGE DETECTION:
- Detect question language automatically
- Respond in the exact same language
- Maintain professional academic tone

FOLLOW-UP DETECTION:
- "that", "it", "this", "summarize", "tell me more", "elaborate", "explain further" → Use conversation history
- New factual questions → Use sources

Remember: You're helping MBE students understand their program requirements clearly and accurately in their preferred language."""
            },
            {
                "role": "user",
                "content": f"""CONVERSATION HISTORY (use for follow-up questions):
{conversation_summary if conversation_summary else "No previous conversation"}

DOCUMENT SOURCES (use for new factual questions):
{context}

CURRENT QUESTION: {query}

Instructions: 
- If this is a follow-up (summarize/elaborate/that/it), answer from conversation history
- If this is a new question, answer from sources with citations
- Always respond in the SAME language as the question
- Always be precise and cite your sources

ANSWER:"""
            }
        ],
        "temperature": 0.1,
        "max_tokens": 2000,
    }
   
    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json=data,
            timeout=60
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"❌ Error connecting to Groq: {str(e)}"

def process_documents_automatically():
    """Auto-process documents on first load"""
    if st.session_state.processed or st.session_state.processing_started:
        return
    
    st.session_state.processing_started = True
    
    available_files = get_files_from_folder()
    if not available_files:
        return
    
    with st.spinner("🔄 Processing documents automatically..."):
        files_data = {}
        all_chunks = []
        all_metadata = []
       
        client = chromadb.Client()
        collection_name = f"docs_{uuid.uuid4().hex[:8]}"
        collection = client.create_collection(
            name=collection_name,
            embedding_function=get_embedding_function()
        )
       
        for filepath in available_files:
            filename = os.path.basename(filepath)
            file_ext = filename.split('.')[-1].lower()
           
            file_hash = get_file_hash(filepath)
            cache_key = f"{file_hash}_{file_ext}"
            cached_data = load_cache(cache_key)
           
            if cached_data:
                file_info = cached_data
                error = None
            else:
                if file_ext == 'pdf':
                    file_info, error = extract_pdf_detailed(filepath)
                elif file_ext in ['docx', 'doc']:
                    file_info, error = extract_docx_detailed(filepath)
                elif file_ext == 'txt':
                    file_info, error = extract_txt_detailed(filepath)
                else:
                    error = "Unsupported file type"
                    file_info = None
                  
                    if error:
                        st.error(f"❌ Error processing {filename}: {error}")
                        continue
                  
                    # Save to cache
                    save_cache(cache_key, file_info)
                    st.success(f"💾 Cached data for: {filename}")
          
            files_data[filename] = file_info
          
            # Add to collection with metadata
            for chunk_obj in file_info['chunks']:
                if isinstance(chunk_obj, dict):
                    all_chunks.append(chunk_obj['content'])
                    all_metadata.append(chunk_obj['metadata'])
                else:
                    all_chunks.append(chunk_obj)
                    all_metadata.append({
                        "source": filename,
                        "page": "N/A",
                        "is_table": "False",
                        "table_number": "N/A"
                    })
          
            progress_bar.progress((idx + 1) / len(available_files))
      
        status_text.text("Building search index...")
      
        if all_chunks:
            batch_size = 500
            for i in range(0, len(all_chunks), batch_size):
                batch = all_chunks[i:i+batch_size]
                metadata_batch = all_metadata[i:i+batch_size]
                collection.add(
                    documents=batch,
                    ids=[f"chunk_{i+j}" for j in range(len(batch))],
                    metadatas=metadata_batch
                )
      
        st.session_state.files_data = files_data
        st.session_state.collection = collection
        st.session_state.processed = True
        st.session_state.processing_started = False  # Reset flag
      
        status_text.empty()
        st.success("✅ Processing completed successfully!")
        st.balloons()
